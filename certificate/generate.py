#!/usr/bin/env python3
"""재직증명서 자동 발급 스크립트 — 공유 양식 기반

양식: template/재직증명서_양식.docx (플레이스홀더 방식)
  {{성명}}    — 직원 이름
  {{소속}}    — 팀명 (없으면 "-")
  {{직위}}    — 직위 (없으면 "-")
  {{주소}}    — 주소 (없으면 "-")
  {{입사일}}  — 입사일 (없으면 "-")
  {{발급일자}} — 오늘 날짜 (자동)

사용법 (CLI):
  python generate.py --create              # 기본 공유 양식 생성
  python generate.py --name 홍길동          # 이름만으로 단건 발급 (정보 없음)
"""

import argparse
import subprocess
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).parent
TEMPLATE_PATH = BASE_DIR / "template" / "재직증명서_양식.docx"
OUTPUT_DIR = BASE_DIR / "output"

# 하위 호환 — 기존 employees/ 디렉토리 (더 이상 사용하지 않음)
EMPLOYEES_DIR = BASE_DIR / "employees"


def _build_replacements(employee: dict) -> dict:
    today = datetime.today().strftime("%Y년 %m월 %d일")
    return {
        "{{성명}}":    employee.get("name", ""),
        "{{소속}}":    employee.get("team", "-") or "-",
        "{{직위}}":    employee.get("position", "-") or "-",
        "{{주소}}":    employee.get("address", "-") or "-",
        "{{입사일}}":  employee.get("joinDate", "-") or "-",
        "{{발급일자}}": today,
    }


def _replace_in_doc(doc: Document, replacements: dict) -> None:
    """DOCX 본문·표 안의 플레이스홀더를 치환한다.

    run이 분할된 경우에도 동작하도록 단락 전체 텍스트를 기준으로 치환 후
    첫 번째 run에 결과를 쓰고 나머지 run을 비운다.
    """
    def _replace_para(para):
        if not any(k in para.text for k in replacements):
            return
        full = "".join(r.text for r in para.runs)
        for key, val in replacements.items():
            full = full.replace(key, val)
        for i, run in enumerate(para.runs):
            run.text = full if i == 0 else ""

    for para in doc.paragraphs:
        _replace_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_para(para)


def _convert_to_pdf(docx_path: Path) -> Path:
    """LibreOffice headless로 PDF 변환."""
    result = subprocess.run(
        [
            "libreoffice", "--headless",
            "--convert-to", "pdf",
            "--outdir", str(docx_path.parent),
            str(docx_path),
        ],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or "libreoffice 변환 실패")
    return docx_path.parent / (docx_path.stem + ".pdf")


def generate_one(employee: dict) -> bool:
    """단건 발급. employee = {name, team, position, address, joinDate}.
    공유 양식이 없으면 자동 생성 후 발급.
    성공 시 True, 실패 시 False.
    """
    name = (employee.get("name") or "").strip()
    if not name:
        print("[SKIP] 이름 없음")
        return False

    if not TEMPLATE_PATH.exists():
        print("[INFO] 공유 양식 없음 — 자동 생성")
        create_base_template()

    today_str = datetime.today().strftime("%Y%m%d")
    stem = f"재직증명서_{name}_{today_str}"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_out = OUTPUT_DIR / f"{stem}.docx"

    doc = Document(TEMPLATE_PATH)
    _replace_in_doc(doc, _build_replacements(employee))
    doc.save(docx_out)
    print(f"[OK]   DOCX: {docx_out.name}")

    try:
        pdf_out = _convert_to_pdf(docx_out)
        print(f"[OK]   PDF : {pdf_out.name}")
    except Exception as e:
        print(f"[WARN] PDF 변환 실패 ({name}): {e}")

    return True


def generate_from_employees(employees: list) -> tuple:
    """Flask API에서 직원 정보 목록을 받아 처리하는 진입점."""
    success, failed = [], []
    for emp in employees:
        name = (emp.get("name") or "").strip()
        if not name:
            continue
        (success if generate_one(emp) else failed).append(name)
    print(f"\n── 완료: 성공 {len(success)}건 / 실패 {len(failed)}건 ──")
    return success, failed


def create_base_template() -> Path:
    """공유 재직증명서 양식을 template/재직증명서_양식.docx 로 생성한다.

    플레이스홀더: {{성명}}, {{소속}}, {{직위}}, {{주소}}, {{입사일}}, {{발급일자}}
    이 파일을 직접 편집해 양식을 꾸밀 수 있다.
    """
    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    if TEMPLATE_PATH.exists():
        print(f"[SKIP] 이미 존재합니다: {TEMPLATE_PATH}")
        return TEMPLATE_PATH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    def add_centered(text, bold=False, size=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        return p

    add_centered("재  직  증  명  서", bold=True, size=22)
    doc.add_paragraph()
    doc.add_paragraph()

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("성    명", "{{성명}}"),
        ("소    속", "{{소속}}"),
        ("직    위", "{{직위}}"),
        ("주    소", "{{주소}}"),
        ("입 사 일", "{{입사일}}"),
    ]
    for i, (label, placeholder) in enumerate(rows_data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = placeholder
    doc.add_paragraph()
    doc.add_paragraph()

    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body.add_run("위 사람은 당사에 재직 중임을 증명합니다.")
    doc.add_paragraph()
    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run("{{발급일자}}")
    doc.add_paragraph()
    doc.add_paragraph()

    add_centered("(주) 이액티브  대표이사  (인)", bold=True)

    doc.save(TEMPLATE_PATH)
    print(f"[OK] 공유 양식 생성: {TEMPLATE_PATH}")
    print("     ※ 이 파일을 직접 편집해 양식을 꾸미세요. 플레이스홀더는 유지하세요.")
    return TEMPLATE_PATH


def main():
    parser = argparse.ArgumentParser(description="재직증명서 자동 발급 (공유 양식 기반)")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--create", action="store_true", help="공유 양식 생성")
    group.add_argument("--name", metavar="이름", help="이름으로 단건 발급 (정보 없이)")
    args = parser.parse_args()

    if args.create:
        create_base_template()
    elif args.name:
        ok = generate_one({"name": args.name})
        sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()
